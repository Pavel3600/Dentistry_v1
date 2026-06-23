from django.views import View
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.urls import reverse_lazy
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse, Http404
from django.core.exceptions import PermissionDenied
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from docx import Document
from django.utils import timezone
import datetime
import asyncio
import json
from .models import (
    UserProfile, Patient, PatientMedicalInfo, Visit, VisitReport,
    Service, Material, MKBSCode, Appointment as ApptModel,
    Clients,
)
from .controllers import (
    PatientController, AppointmentController,
    MedicalRecordController, MKBSController,
    ReferralController, StudyController,
    WorkOrderController, ClientController,
    VisitController, ServiceController, MaterialController,
    AppointmentLogController, PatientMedicalInfoController,
)
from .forms import (PatientForm, AppointmentForm, VisitForm, ProcedureForm,
                    ReferralForm, ServiceForm, MaterialForm, MkbCodeForm,
                    VisitSearchForm, RegistrationForm, AdminAppointmentForm,
                    DoctorForm, ManagerForm, PatientMedicalInfoForm, VisitReportForm)
from django.contrib.auth.models import User
from django.contrib.auth import views as auth_views, login as auth_login, logout
from .services.fastapi_client import fastapi_client
from .utils import is_fastapi_available


# ==================== МИКСИНЫ ДЛЯ ПРОВЕРКИ РОЛЕЙ ====================

class RoleRequiredMixin(LoginRequiredMixin, UserPassesTestMixin):
    """Базовый миксин для проверки ролей"""
    login_url = 'login'

    def test_func(self):
        return hasattr(self.request.user, 'profile') and self.request.user.profile.role in self.allowed_roles


class AdminRequiredMixin(RoleRequiredMixin):
    """Миксин для доступа только администраторам"""
    allowed_roles = ['admin']


class ManagerRequiredMixin(RoleRequiredMixin):
    """Миксин для доступа менеджерам и администраторам"""
    allowed_roles = ['manager', 'admin']


class DentistRequiredMixin(RoleRequiredMixin):
    """Миксин для доступа только стоматологам"""
    allowed_roles = ['dentist']


class PatientRequiredMixin(RoleRequiredMixin):
    """Миксин для доступа только пациентам"""
    allowed_roles = ['patient']


class StaffRequiredMixin(RoleRequiredMixin):
    """Доступ персоналу клиники: менеджер, стоматолог, администратор."""
    allowed_roles = ['manager', 'dentist', 'admin']


# ==================== УТИЛИТЫ ====================

class AsyncHelper:
    """Утилита для работы с асинхронными вызовами"""

    @staticmethod
    def run_async(coro):
        """Запускает асинхронную функцию синхронно"""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(coro)
        finally:
            loop.close()


class PatientHelper:
    """Утилита для работы с пациентами"""

    @staticmethod
    def get_patient_for_user(user):
        """Получает профиль пациента для пользователя через API."""
        try:
            patients = PatientController.get_all(size=500)
            for p in patients:
                if p.get('user_id') == user.id:
                    return p
        except Exception:
            pass
        return None

    @staticmethod
    def ensure_client_for_patient(full_name):
        """Создаёт клиента(role=patient) через API и возвращает его id."""
        from django.utils.text import slugify
        base = slugify(full_name, allow_unicode=False) or 'patient'
        login = f"{base}-{timezone.now().strftime('%H%M%S%f')[:9]}"
        try:
            client = ClientController.create({'login': login, 'password': 'pat!2024'})
            return client.get('id')
        except Exception:
            return None


# ==================== БАЗОВЫЕ КЛАССЫ ДЛЯ CRUD ====================


# ==================== КЛАССЫ И ПРОВЕРКИ ====================

class CustomLoginView(auth_views.LoginView):
    """Кастомное представление входа с проверкой FastAPI"""
    template_name = 'registration/login.html'

    def dispatch(self, request, *args, **kwargs):
        if not is_fastapi_available():
            messages.error(request, "⛔ Ошибка системы: Сервис FastAPI недоступен. Авторизация временно запрещена.")
            return self.render_to_response(self.get_context_data(form=self.get_form()))
        return super().dispatch(request, *args, **kwargs)


# ==================== ОБЩИЕ ФУНКЦИИ ====================

@login_required
def index(request):
    """Главная страница с редиректом по роли"""
    role_redirects = {
        'dentist': 'app:doctor_dashboard',
        'manager': 'app:manager_dashboard',
        'admin': 'app:admin_dashboard',
    }

    user_role = getattr(request.user.profile, 'role', None) if hasattr(request.user, 'profile') else None
    redirect_url = role_redirects.get(user_role, 'app:admin_dashboard')
    return redirect(redirect_url)


@login_required
def logout_view(request):
    logout(request)
    return redirect('login')


def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Регистрация успешна! Теперь вы можете войти.')
            return redirect('login')
    else:
        form = RegistrationForm()
    return render(request, 'registration/register.html', {'form': form})


# ==================== АДМИНИСТРАТОР - CBV ДЛЯ CRUD ====================

class AdminDashboardView(AdminRequiredMixin, View):
    template_name = 'admin_panel/admin_dashboard.html'

    def get(self, request):
        return render(request, self.template_name, {'title': 'Панель Администратора'})


class AdminAppointmentCreateView(AdminRequiredMixin, View):
    template_name = 'admin_panel/appointment_create.html'

    def get(self, request):
        form = AdminAppointmentForm()
        return render(request, self.template_name, {'form': form})

    def post(self, request):
        form = AdminAppointmentForm(request.POST)
        if form.is_valid():
            try:
                data = form.cleaned_data
                if hasattr(data.get('datetime'), 'isoformat'):
                    dt = data['datetime']
                    if hasattr(dt, 'tzinfo') and dt.tzinfo:
                        dt = dt.replace(tzinfo=None)
                    data['datetime'] = dt.isoformat()
                result = AppointmentController.create(data)
                messages.success(request, f"Запись успешно создана! ID: {result.get('id')}")
                return redirect('app:appointment_list')
            except Exception as e:
                messages.error(request, f"Ошибка при создании: {e}")
        else:
            messages.error(request, "Проверьте правильность заполнения полей.")
        return render(request, self.template_name, {'form': form})


class ServiceListView(AdminRequiredMixin, View):
    template_name = 'admin_panel/service_list.html'

    def get(self, request):
        try:
            services = ServiceController.get_all()
        except Exception:
            services = []
        return render(request, self.template_name, {'services': services})


class ServiceCreateView(AdminRequiredMixin, View):
    template_name = 'admin_panel/service_form.html'

    def get(self, request):
        return render(request, self.template_name, {'form': ServiceForm()})

    def post(self, request):
        form = ServiceForm(request.POST)
        if form.is_valid():
            try:
                ServiceController.create(form.cleaned_data)
                messages.success(request, 'Услуга успешно создана')
                return redirect('app:service_list')
            except Exception as e:
                messages.error(request, f'Ошибка FastAPI: {e}')
        return render(request, self.template_name, {'form': form})


class ServiceUpdateView(AdminRequiredMixin, View):
    template_name = 'admin_panel/service_form.html'

    def get(self, request, pk):
        try:
            service = ServiceController.get_by_id(pk)
        except Exception:
            raise Http404
        form = ServiceForm(initial=service)
        return render(request, self.template_name, {'form': form, 'object': service})

    def post(self, request, pk):
        form = ServiceForm(request.POST)
        if form.is_valid():
            try:
                ServiceController.update(pk, form.cleaned_data)
                messages.success(request, 'Услуга успешно обновлена')
                return redirect('app:service_list')
            except Exception as e:
                messages.error(request, f'Ошибка FastAPI: {e}')
        return render(request, self.template_name, {'form': form})


class ServiceDeleteView(AdminRequiredMixin, View):

    def post(self, request, pk):
        try:
            ServiceController.delete(pk)
            messages.success(request, 'Услуга удалена')
        except Exception as e:
            messages.error(request, f'Ошибка FastAPI: {e}')
        return redirect('app:service_list')


class MaterialListView(AdminRequiredMixin, View):
    template_name = 'admin_panel/material_list.html'

    def get(self, request):
        try:
            materials = MaterialController.get_all()
        except Exception:
            materials = []
        return render(request, self.template_name, {'materials': materials})


class MaterialCreateView(AdminRequiredMixin, View):
    template_name = 'admin_panel/material_form.html'

    def get(self, request):
        return render(request, self.template_name, {'form': MaterialForm()})

    def post(self, request):
        form = MaterialForm(request.POST)
        if form.is_valid():
            try:
                MaterialController.create(form.cleaned_data)
                messages.success(request, 'Материал успешно создан')
                return redirect('app:material_list')
            except Exception as e:
                messages.error(request, f'Ошибка FastAPI: {e}')
        return render(request, self.template_name, {'form': form})


class MaterialUpdateView(AdminRequiredMixin, View):
    template_name = 'admin_panel/material_form.html'

    def get(self, request, pk):
        try:
            material = MaterialController.get_by_id(pk)
        except Exception:
            raise Http404
        form = MaterialForm(initial=material)
        return render(request, self.template_name, {'form': form, 'object': material})

    def post(self, request, pk):
        form = MaterialForm(request.POST)
        if form.is_valid():
            try:
                MaterialController.update(pk, form.cleaned_data)
                messages.success(request, 'Материал успешно обновлён')
                return redirect('app:material_list')
            except Exception as e:
                messages.error(request, f'Ошибка FastAPI: {e}')
        return render(request, self.template_name, {'form': form})


class MaterialDeleteView(AdminRequiredMixin, View):

    def post(self, request, pk):
        try:
            MaterialController.delete(pk)
            messages.success(request, 'Материал удалён')
        except Exception as e:
            messages.error(request, f'Ошибка FastAPI: {e}')
        return redirect('app:material_list')


class MkbCodeListView(AdminRequiredMixin, View):
    template_name = 'admin_panel/mkb_list.html'

    def get(self, request):
        try:
            codes = MKBSController.get_all()
        except Exception:
            codes = []
        return render(request, self.template_name, {'codes': codes})


class MkbCodeCreateView(AdminRequiredMixin, View):
    template_name = 'admin_panel/mkb_form.html'
    success_url = reverse_lazy('app:mkb_list')

    def get(self, request):
        return render(request, self.template_name, {'form': MkbCodeForm()})

    def post(self, request):
        form = MkbCodeForm(request.POST)
        if form.is_valid():
            try:
                MKBSController.create(form.cleaned_data)
                messages.success(request, 'Код МКБ успешно создан')
                return redirect(self.success_url)
            except Exception as e:
                messages.error(request, f'Ошибка FastAPI: {e}')
        return render(request, self.template_name, {'form': form})


class MkbCodeUpdateView(AdminRequiredMixin, View):
    template_name = 'admin_panel/mkb_form.html'
    success_url = reverse_lazy('app:mkb_list')

    def get(self, request, pk):
        try:
            code = MKBSController.get_by_id(pk)
        except Exception:
            raise Http404
        form = MkbCodeForm(initial=code)
        return render(request, self.template_name, {'form': form, 'code': code})

    def post(self, request, pk):
        form = MkbCodeForm(request.POST)
        if form.is_valid():
            try:
                MKBSController.update(pk, form.cleaned_data)
                messages.success(request, 'Код МКБ обновлён')
                return redirect(self.success_url)
            except Exception as e:
                messages.error(request, f'Ошибка FastAPI: {e}')
        return render(request, self.template_name, {'form': form})


class MkbCodeDeleteView(AdminRequiredMixin, View):
    success_url = reverse_lazy('app:mkb_list')

    def post(self, request, pk):
        try:
            MKBSController.delete(pk)
            messages.success(request, f'Код #{pk} удалён')
        except Exception as e:
            messages.error(request, f'Ошибка FastAPI: {e}')
        return redirect(self.success_url)


# ==================== МЕНЕДЖЕР ====================

class ManagerDashboardView(ManagerRequiredMixin, View):
    template_name = 'manager/manager_dashboard.html'

    def get(self, request):
        today = timezone.now().date()

        try:
            patients = PatientController.get_all(size=200)
            appointments_all = AppointmentController.get_all(size=500)
        except Exception:
            patients = []
            appointments_all = []

        today_str = today.isoformat()
        month_ago_str = (today - datetime.timedelta(days=30)).isoformat()

        scheduled_today = sum(
            1 for a in appointments_all
            if a.get('status') == 'scheduled' and str(a.get('datetime', ''))[:10] == today_str
        )
        upcoming = sum(
            1 for a in appointments_all
            if a.get('status') == 'scheduled' and str(a.get('datetime', ''))[:10] >= today_str
        )
        completed_month = sum(
            1 for a in appointments_all
            if a.get('status') == 'completed' and str(a.get('datetime', ''))[:10] >= month_ago_str
        )

        context = {
            'total_patients': len(patients),
            'scheduled_today': scheduled_today,
            'upcoming_appointments': upcoming,
            'completed_this_month': completed_month,
            'today_appointments': self._get_today_appointments(appointments_all, patients, today_str),
            'new_patients': patients[-5:],
            'popular_doctors': [],
        }
        return render(request, self.template_name, context)

    def _get_today_appointments(self, appointments_all, patients, today_str):
        patient_map = {p['id']: p.get('full_name', '') for p in patients}
        result = []
        for a in appointments_all:
            if a.get('status') == 'scheduled' and str(a.get('datetime', ''))[:10] == today_str:
                result.append({
                    'id': a['id'],
                    'datetime': a.get('datetime'),
                    'patient_name': patient_map.get(a.get('patient_id'), f"Пациент #{a.get('patient_id')}"),
                    'doctor_name': f"Врач #{a.get('doctor_id')}",
                    'status': a.get('status'),
                })
                if len(result) >= 10:
                    break
        return result


class PatientListView(ManagerRequiredMixin, View):
    template_name = 'manager/patient_list.html'

    def get(self, request):
        page = int(request.GET.get('page', 1))
        patients = list(Patient.objects.all()[(page - 1) * 50: page * 50])
        return render(request, self.template_name, {'patients': patients, 'page': page})


class PatientCreateView(ManagerRequiredMixin, View):
    template_name = 'manager/patient_form.html'
    success_url = reverse_lazy('app:patient_list')

    def get(self, request):
        return render(request, self.template_name, {'form': PatientForm()})

    def post(self, request):
        form = PatientForm(request.POST)
        if form.is_valid():
            cd = form.cleaned_data
            from django.utils.text import slugify
            base = slugify(cd['full_name'], allow_unicode=False) or 'patient'
            login = f"{base}-{timezone.now().strftime('%H%M%S%f')[:9]}"
            client = Clients.objects.create(login=login, role='patient')
            patient = Patient.objects.create(
                full_name=cd['full_name'],
                birth_date=cd['birth_date'],
                gender=cd['gender'],
                phone=cd['phone'],
                address=cd.get('address') or '',
                user_id=client.id,
            )
            messages.success(request, 'Пациент успешно добавлен')
            return redirect(self.success_url)
        return render(request, self.template_name, {'form': form})


class AppointmentCreateView(ManagerRequiredMixin, View):
    template_name = 'manager/appointment_form.html'
    success_url = reverse_lazy('app:appointment_list')

    def _get_context(self, form):
        # Пациенты — только из локальной Django-БД (чтобы pk совпадал при сохранении)
        patients = list(Patient.objects.all().order_by('full_name'))
        # Врачи — Django-пользователи с ролью dentist
        from app.models import UserProfile as UP
        doctor_profiles = UP.objects.filter(role='dentist').select_related('user').order_by(
            'user__last_name', 'user__first_name'
        )
        doctors = [
            {'id': p.user.pk, 'login': p.user.get_full_name() or p.user.username}
            for p in doctor_profiles
        ]
        return {'form': form, 'patients': patients, 'doctors': doctors}

    def get(self, request, *args, **kwargs):
        form = AppointmentForm(initial={'patient_id': request.GET.get('patient_id')})
        return render(request, self.template_name, self._get_context(form))

    def post(self, request, *args, **kwargs):
        form = AppointmentForm(request.POST)
        if form.is_valid():
            try:
                data = form.cleaned_data
                if hasattr(data.get('datetime'), 'isoformat'):
                    dt = data['datetime']
                    if hasattr(dt, 'tzinfo') and dt.tzinfo:
                        dt = dt.replace(tzinfo=None)
                    data['datetime'] = dt.isoformat()
                patient_obj = Patient.objects.get(pk=data['patient_id'])
                doctor_obj = User.objects.get(pk=data['doctor_id'])
                appt = ApptModel.objects.create(
                    patient=patient_obj, doctor=doctor_obj,
                    datetime=data['datetime'], status='scheduled',
                )
                messages.success(request, f'Запись успешно создана! ID: {appt.id}')
                return redirect(self.success_url)
            except Exception as e:
                messages.error(request, f'Ошибка: {e}')
        else:
            messages.error(request, 'Пожалуйста, исправьте ошибки в форме.')
        return render(request, self.template_name, self._get_context(form))


class PatientUpdateView(ManagerRequiredMixin, View):
    template_name = 'manager/patient_form.html'
    success_url = reverse_lazy('app:patient_list')

    def get(self, request, pk):
        from django.http import Http404
        try:
            patient = Patient.objects.get(pk=pk)
        except Patient.DoesNotExist:
            raise Http404
        form = PatientForm(initial={'full_name': patient.full_name, 'birth_date': patient.birth_date,
                                    'gender': patient.gender, 'phone': patient.phone, 'address': patient.address})
        return render(request, self.template_name, {'form': form, 'patient': patient})

    def post(self, request, pk):
        from django.http import Http404
        try:
            patient = Patient.objects.get(pk=pk)
        except Patient.DoesNotExist:
            raise Http404
        form = PatientForm(request.POST)
        if form.is_valid():
            cd = form.cleaned_data
            patient.full_name = cd['full_name']
            patient.birth_date = cd['birth_date']
            patient.gender = cd['gender']
            patient.phone = cd['phone']
            patient.address = cd.get('address', '')
            patient.save()
            messages.success(request, 'Данные пациента обновлены')
            return redirect(self.success_url)
        return render(request, self.template_name, {'form': form})


STATUS_CHOICES = [
    ('scheduled', 'Запланирован'),
    ('completed', 'Завершён'),
    ('cancelled', 'Отменён'),
]


class AppointmentListView(ManagerRequiredMixin, View):
    template_name = 'manager/appointment_list.html'

    def get(self, request):
        page = int(request.GET.get('page', 1))
        status_filter = request.GET.get('status', '')
        qs = ApptModel.objects.select_related('patient', 'doctor').order_by('-datetime')
        if status_filter:
            qs = qs.filter(status=status_filter)
        appointments = list(qs[(page - 1) * 50: page * 50])
        return render(request, self.template_name, {
            'appointments': appointments,
            'status_choices': STATUS_CHOICES,
            'current_status': status_filter,
            'page': page,
        })


class AppointmentChangeStatusView(ManagerRequiredMixin, View):
    """Сменить статус записи."""

    def post(self, request, pk):
        from django.http import Http404
        new_status = request.POST.get('status')
        valid = dict(STATUS_CHOICES)
        if new_status not in valid:
            messages.error(request, 'Недопустимый статус.')
            return redirect('app:appointment_list')
        try:
            appointment = ApptModel.objects.get(pk=pk)
        except ApptModel.DoesNotExist:
            raise Http404
        old_status = appointment.status
        if new_status == 'completed' and not hasattr(appointment, 'visit'):
            return redirect('app:start_visit', pk=pk)
        appointment.status = new_status
        appointment.save()
        from .models import AppointmentLog
        AppointmentLog.objects.create(
            appointment=appointment,
            changed_by=request.user,
            old_status=old_status,
            new_status=new_status,
        )
        messages.success(request, f'Запись #{pk}: статус «{valid[new_status]}».')
        return redirect('app:appointment_list')


class AppointmentDeleteView(ManagerRequiredMixin, View):
    """Удалить запись на приём."""

    def post(self, request, pk):
        ApptModel.objects.filter(pk=pk).delete()
        messages.success(request, f'Запись #{pk} удалена.')
        return redirect('app:appointment_list')


class AppointmentCancelView(ManagerRequiredMixin, View):

    def post(self, request, pk):
        from django.http import Http404
        try:
            appointment = ApptModel.objects.get(pk=pk)
        except ApptModel.DoesNotExist:
            raise Http404
        old_status = appointment.status
        appointment.status = 'cancelled'
        appointment.save()
        from .models import AppointmentLog
        AppointmentLog.objects.create(
            appointment=appointment,
            changed_by=request.user,
            old_status=old_status,
            new_status='cancelled',
        )
        messages.success(request, f'Запись #{pk} отменена')
        return redirect('app:appointment_list')


# ==================== УПРАВЛЕНИЕ ВРАЧАМИ ====================

class DoctorListView(ManagerRequiredMixin, View):
    """Список врачей-стоматологов."""
    template_name = 'manager/doctor_list.html'

    def get(self, request):
        doctors = UserProfile.objects.filter(role='dentist').select_related('user').order_by(
            'user__last_name', 'user__first_name'
        )
        return render(request, self.template_name, {'doctors': doctors})


class DoctorCreateView(ManagerRequiredMixin, View):
    """Удобное добавление врача."""
    template_name = 'manager/doctor_form.html'
    success_url = reverse_lazy('app:doctor_list')

    def get(self, request):
        return render(request, self.template_name, {'form': DoctorForm()})

    def post(self, request):
        form = DoctorForm(request.POST)
        if form.is_valid():
            doctor = form.save()
            name = doctor.get_full_name() or doctor.username
            password = form.cleaned_data['password']
            messages.success(
                request,
                f'Врач {name} успешно добавлен. '
                f'Логин: {doctor.username} | Пароль: {password}'
            )
            return redirect(self.success_url)
        messages.error(request, 'Проверьте правильность заполнения полей.')
        return render(request, self.template_name, {'form': form})


class ManagerListView(AdminRequiredMixin, View):
    """Список менеджеров регистратуры — только для администратора."""
    template_name = 'manager/manager_list.html'

    def get(self, request):
        managers = UserProfile.objects.filter(role='manager').select_related('user').order_by(
            'user__last_name', 'user__first_name'
        )
        return render(request, self.template_name, {'managers': managers})


class ManagerCreateView(AdminRequiredMixin, View):
    """Создание менеджера регистратуры — только для администратора."""
    template_name = 'manager/manager_form.html'
    success_url = reverse_lazy('app:manager_list')

    def get(self, request):
        return render(request, self.template_name, {'form': ManagerForm()})

    def post(self, request):
        form = ManagerForm(request.POST)
        if form.is_valid():
            mgr = form.save()
            name = mgr.get_full_name() or mgr.username
            password = form.cleaned_data['password']
            messages.success(
                request,
                f'Менеджер {name} успешно добавлен. '
                f'Логин: {mgr.username} | Пароль: {password}'
            )
            return redirect(self.success_url)
        messages.error(request, 'Проверьте правильность заполнения полей.')
        return render(request, self.template_name, {'form': form})


# ==================== СТОМАТОЛОГ ====================

class DoctorDashboardView(DentistRequiredMixin, View):
    template_name = 'doctor/doctor_dashboard.html'

    def get(self, request):
        today = timezone.now().date()
        doctor_profile = self._get_doctor_profile(request.user)

        context = {
            'today_appointments': self._get_today_appointments(doctor_profile, today),
            'upcoming_appointments': self._get_upcoming_appointments(doctor_profile, today),
            'recent_visits': self._get_recent_visits(doctor_profile),
            'stats': self._get_stats(doctor_profile, today),
            'week_schedule': self._get_week_schedule(doctor_profile, today),
            'today': today,
            'doctor_name': request.user.get_full_name() or request.user.username,
        }
        return render(request, self.template_name, context)

    def _get_doctor_profile(self, user):
        try:
            return user.profile
        except UserProfile.DoesNotExist:
            return UserProfile.objects.create(user=user, role='dentist')

    def _get_today_appointments(self, doctor_profile, today):
        today_str = today.isoformat()
        try:
            appts = AppointmentController.get_all(doctor_id=doctor_profile.id, size=200)
            return [a for a in appts if a.get('status') == 'scheduled'
                    and str(a.get('datetime', ''))[:10] == today_str]
        except Exception:
            return []

    def _get_upcoming_appointments(self, doctor_profile, today):
        today_str = today.isoformat()
        next_week_str = (today + datetime.timedelta(days=7)).isoformat()
        try:
            appts = AppointmentController.get_all(doctor_id=doctor_profile.id, size=200)
            return [a for a in appts
                    if a.get('status') == 'scheduled'
                    and today_str < str(a.get('datetime', ''))[:10] <= next_week_str][:10]
        except Exception:
            return []

    def _get_recent_visits(self, doctor_profile):
        try:
            return VisitController.get_all(doctor_id=doctor_profile.id, size=5)
        except Exception:
            return []

    def _get_stats(self, doctor_profile, today):
        month_ago_str = (today - datetime.timedelta(days=30)).isoformat()
        try:
            visits = VisitController.get_all(doctor_id=doctor_profile.id, size=500)
            recent = [v for v in visits if str(v.get('visit_date', ''))[:10] >= month_ago_str]
            patient_ids = {v['patient_id'] for v in recent}
        except Exception:
            recent = []
            patient_ids = set()
        return {
            'total_patients': len(patient_ids),
            'total_visits': len(recent),
            'total_procedures': 0,
        }

    def _get_week_schedule(self, doctor_profile, today):
        week_schedule = []
        day_names = ['Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс']
        try:
            appts = AppointmentController.get_all(doctor_id=doctor_profile.id, size=500)
        except Exception:
            appts = []
        for i in range(7):
            day = today + datetime.timedelta(days=i)
            day_str = day.isoformat()
            count = sum(1 for a in appts
                        if a.get('status') == 'scheduled' and str(a.get('datetime', ''))[:10] == day_str)
            week_schedule.append({'date': day, 'day_name': day_names[day.weekday()], 'count': count})
        return week_schedule


class AppointmentDetailView(StaffRequiredMixin, View):
    template_name = 'doctor/appointment_detail.html'

    def get(self, request, pk):
        from django.http import Http404
        try:
            appointment = ApptModel.objects.select_related('patient', 'doctor').get(pk=pk)
        except ApptModel.DoesNotExist:
            raise Http404
        visit = getattr(appointment, 'visit', None)
        patient = appointment.patient
        diagnoses = list(MKBSCode.objects.filter(is_active=True))
        reports = list(VisitReport.objects.filter(visit=visit)) if visit else []
        return render(request, self.template_name, {
            'appointment': appointment,
            'visit': visit,
            'patient': patient,
            'reports': reports,
            'diagnoses': diagnoses,
        })


class StartVisitView(StaffRequiredMixin, View):
    """Завершение приёма и создание медкарты."""
    template_name = 'doctor/complete_visit.html'

    def get(self, request, pk):
        from django.http import Http404
        try:
            appointment = ApptModel.objects.select_related('patient', 'doctor').get(pk=pk)
        except ApptModel.DoesNotExist:
            raise Http404
        if hasattr(appointment, 'visit'):
            return redirect('app:appointment_detail', pk=pk)
        diagnoses = list(MKBSCode.objects.filter(is_active=True))
        return render(request, self.template_name, {
            'appointment': appointment,
            'patient': appointment.patient,
            'diagnoses': diagnoses,
        })

    def post(self, request, pk):
        from django.http import Http404
        try:
            appointment = ApptModel.objects.select_related('patient', 'doctor').get(pk=pk)
        except ApptModel.DoesNotExist:
            raise Http404

        diagnosis_id = request.POST.get('diagnosis') or None
        diagnosis = None
        if diagnosis_id:
            diagnosis = MKBSCode.objects.filter(pk=diagnosis_id).first()

        # Убеждаемся что запись Clients для врача существует
        Clients.objects.get_or_create(
            login=appointment.doctor.username,
            defaults={'role': 'dentist'},
        )
        if not hasattr(appointment, 'visit'):
            patient = appointment.patient
            if not patient:
                from django.http import Http404
                raise Http404
            visit = Visit.objects.create(
                appointment=appointment,
                patient=patient,
                doctor=appointment.doctor,
                anamnesis=request.POST.get('anamnesis', ''),
                examination_results=request.POST.get('examination_results', ''),
                diagnosis=diagnosis,
                treatment_plan=request.POST.get('treatment_plan', ''),
                prescription=request.POST.get('prescription', ''),
                tooth_formula=request.POST.get('tooth_formula', ''),
            )
        old_status = appointment.status
        appointment.status = 'completed'
        appointment.save()
        from .models import AppointmentLog
        AppointmentLog.objects.create(
            appointment=appointment,
            changed_by=request.user,
            old_status=old_status,
            new_status='completed',
        )
        messages.success(request, 'Приём завершён, медкарта сохранена.')
        return redirect('app:appointment_detail', pk=pk)


class ProcedureCreateView(DentistRequiredMixin, View):
    template_name = 'doctor/procedure_form.html'

    def get(self, request, visit_pk):
        from django.http import Http404
        try:
            visit = Visit.objects.select_related('appointment').get(pk=visit_pk)
        except Visit.DoesNotExist:
            raise Http404
        services = list(Service.objects.all())
        return render(request, self.template_name, {
            'form': ProcedureForm(), 'visit_pk': visit_pk,
            'services': services, 'appointment_id': visit.appointment_id,
        })

    def post(self, request, visit_pk):
        from django.http import Http404
        try:
            visit = Visit.objects.select_related('appointment').get(pk=visit_pk)
        except Visit.DoesNotExist:
            raise Http404
        form = ProcedureForm(request.POST)
        if form.is_valid():
            from .models import Procedure as ProcModel
            ProcModel.objects.create(
                visit=visit,
                service=form.cleaned_data['service'],
                quantity=form.cleaned_data['quantity'],
            )
            messages.success(request, 'Процедура добавлена')
            return redirect('app:appointment_detail', pk=visit.appointment_id)
        services = list(Service.objects.all())
        return render(request, self.template_name, {
            'form': form, 'visit_pk': visit_pk,
            'services': services, 'appointment_id': visit.appointment_id,
        })


class ReferralCreateView(DentistRequiredMixin, View):
    template_name = 'doctor/referral_form.html'

    def get(self, request, visit_pk):
        try:
            visit = VisitController.get_by_id(visit_pk)
        except Exception:
            from django.http import Http404
            raise Http404
        return render(request, self.template_name, {'form': ReferralForm(), 'visit': visit})

    def post(self, request, visit_pk):
        try:
            visit = VisitController.get_by_id(visit_pk)
        except Exception:
            from django.http import Http404
            raise Http404
        form = ReferralForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            data['patient_id'] = visit.get('patient_id')
            data['doctor_id'] = visit.get('doctor_id')
            try:
                ReferralController.create(data)
                messages.success(request, 'Направление создано')
                return redirect('app:appointment_detail', pk=visit.get('appointment_id'))
            except Exception as e:
                messages.error(request, f'Ошибка API: {e}')
        return render(request, self.template_name, {'form': form, 'visit': visit})


class CreateExtractWordView(LoginRequiredMixin, View):
    """Выписка в Word. Доступна персоналу (любая выписка) и пациенту (только своя)."""
    login_url = 'login'

    def get(self, request, visit_pk):
        from django.http import Http404
        try:
            visit = Visit.objects.select_related('patient', 'doctor', 'diagnosis').get(pk=visit_pk)
        except Visit.DoesNotExist:
            raise Http404

        role = getattr(getattr(request.user, 'profile', None), 'role', 'patient')
        if role not in ('admin', 'manager', 'dentist'):
            if not visit.patient or visit.patient.user_id != request.user.id:
                raise PermissionDenied('Доступ к чужой выписке запрещён.')

        doc = self._create_document(visit)
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="extract_{visit.pk}.docx"'
        doc.save(response)
        return response

    def _create_document(self, visit):
        doc = Document()
        doc.add_heading('Медицинская карта стоматологического больного (форма 043/у)', 0)
        self._add_patient_info(doc, visit)
        self._add_medical_info(doc, visit)
        self._add_procedures(doc, visit)
        return doc

    def _add_patient_info(self, doc, visit):
        p = visit.patient
        doc.add_paragraph(f'Пациент: {p.full_name if p else "—"}')
        doc.add_paragraph(f'Дата рождения: {str(p.birth_date)[:10] if p else "—"}')
        doc.add_paragraph(f'Дата приёма: {str(visit.visit_date)[:10]}')
        doc.add_paragraph(f'Врач: {visit.doctor.get_full_name() if visit.doctor else "—"}')

    def _add_medical_info(self, doc, visit):
        doc.add_paragraph(f'Анамнез: {visit.anamnesis or "—"}')
        doc.add_paragraph(f'Результаты осмотра: {visit.examination_results or "—"}')
        diag = visit.diagnosis.name if visit.diagnosis else '—'
        doc.add_paragraph(f'Диагноз: {diag}')
        doc.add_paragraph(f'Лечение: {visit.treatment_plan or "—"}')
        doc.add_paragraph(f'Назначения: {visit.prescription or "—"}')
        doc.add_paragraph(f'Зубная формула: {visit.tooth_formula or "—"}')

    def _add_procedures(self, doc, visit):
        doc.add_paragraph('--- Проведённые манипуляции ---')
        for proc in visit.procedures.all():
            doc.add_paragraph(f'- {proc.service.name} x{proc.quantity} = {proc.total_cost} руб.')


class PatientHistoryView(DentistRequiredMixin, View):
    template_name = 'doctor/patient_history.html'

    def get(self, request, patient_pk):
        from django.http import Http404
        try:
            patient = Patient.objects.get(pk=patient_pk)
        except Patient.DoesNotExist:
            raise Http404
        visits = list(Visit.objects.filter(patient=patient).order_by('-visit_date'))
        medical_info, _ = PatientMedicalInfo.objects.get_or_create(patient=patient)
        return render(request, self.template_name, {
            'patient': patient,
            'visits': visits,
            'medical_info': medical_info,
        })


class PatientMedicalInfoView(DentistRequiredMixin, View):
    """Просмотр и редактирование мед. предупреждений (аллергии и пр.) о пациенте."""
    template_name = 'doctor/medical_info_form.html'

    def _get_patient(self, patient_pk):
        from django.http import Http404
        try:
            return Patient.objects.get(pk=patient_pk)
        except Patient.DoesNotExist:
            raise Http404

    def get(self, request, patient_pk):
        patient = self._get_patient(patient_pk)
        info, _ = PatientMedicalInfo.objects.get_or_create(patient=patient)
        form = PatientMedicalInfoForm(initial={
            'allergies': info.allergies, 'chronic_conditions': info.chronic_conditions,
            'contraindications': info.contraindications, 'blood_type': info.blood_type,
            'notes': info.notes,
        })
        return render(request, self.template_name, {'form': form, 'patient': patient, 'info': info})

    def post(self, request, patient_pk):
        patient = self._get_patient(patient_pk)
        form = PatientMedicalInfoForm(request.POST)
        if form.is_valid():
            cd = form.cleaned_data
            info, _ = PatientMedicalInfo.objects.get_or_create(patient=patient)
            for field, value in cd.items():
                setattr(info, field, value)
            info.save()
            messages.success(request, 'Мед. информация пациента сохранена')
            return redirect('app:patient_history', patient_pk=patient_pk)
        info, _ = PatientMedicalInfo.objects.get_or_create(patient=patient)
        return render(request, self.template_name, {'form': form, 'patient': patient, 'info': info})


class VisitReportCreateView(StaffRequiredMixin, View):
    """Отчёт после приёма."""
    template_name = 'doctor/report_form.html'

    def _get_visit(self, visit_pk):
        from django.http import Http404
        try:
            return Visit.objects.get(pk=visit_pk)
        except Visit.DoesNotExist:
            raise Http404

    def get(self, request, visit_pk):
        visit = self._get_visit(visit_pk)
        reports = list(VisitReport.objects.filter(visit=visit))
        return render(request, self.template_name, {'form': VisitReportForm(), 'visit': visit, 'reports': reports})

    def post(self, request, visit_pk):
        visit = self._get_visit(visit_pk)
        form = VisitReportForm(request.POST)
        if form.is_valid():
            cd = form.cleaned_data
            if not cd.get('summary'):
                form.add_error('summary', 'Обязательное поле.')
            else:
                VisitReport.objects.update_or_create(
                    visit=visit,
                    defaults={
                        'author': request.user,
                        'title': cd.get('title', ''),
                        'summary': cd['summary'],
                        'recommendations': cd.get('recommendations', ''),
                        'complications': cd.get('complications', ''),
                    }
                )
                messages.success(request, 'Отчёт сохранён')
                return redirect('app:visit_report', visit_pk=visit_pk)
        reports = list(VisitReport.objects.filter(visit=visit))
        return render(request, self.template_name, {'form': form, 'visit': visit, 'reports': reports})


class VisitReportUpdateView(StaffRequiredMixin, View):
    """Редактирование отчёта о визите."""
    template_name = 'doctor/report_form.html'

    def get(self, request, pk):
        from django.http import Http404
        try:
            report = VisitReport.objects.get(pk=pk)
        except VisitReport.DoesNotExist:
            raise Http404
        form = VisitReportForm(initial={
            'title': report.title, 'summary': report.summary,
            'recommendations': report.recommendations, 'complications': report.complications,
        })
        return render(request, self.template_name, {
            'form': form, 'report': report,
            'editing': report, 'visit': report.visit,
        })

    def post(self, request, pk):
        from django.http import Http404
        try:
            report = VisitReport.objects.get(pk=pk)
        except VisitReport.DoesNotExist:
            raise Http404
        form = VisitReportForm(request.POST)
        if form.is_valid():
            cd = form.cleaned_data
            report.title = cd.get('title', '')
            report.summary = cd['summary']
            report.recommendations = cd.get('recommendations', '')
            report.complications = cd.get('complications', '')
            report.save()
            messages.success(request, 'Отчёт обновлён')
            return redirect('app:visit_report', visit_pk=report.visit_id)
        return render(request, self.template_name, {
            'form': form, 'report': report,
            'editing': report, 'visit': report.visit,
        })


class VisitReportDeleteView(StaffRequiredMixin, View):
    """Удаление отчёта о приёме."""

    def post(self, request, pk):
        VisitReport.objects.filter(pk=pk).delete()
        messages.success(request, 'Отчёт удалён')
        return redirect('app:appointment_list')


class SearchByDateView(DentistRequiredMixin, View):
    template_name = 'doctor/search_results.html'

    def get(self, request):
        form = VisitSearchForm(request.GET)
        visits = self._filter_visits(form, request.GET)

        try:
            services = ServiceController.get_all()
        except Exception:
            services = []

        return render(request, self.template_name, {
            'visits': visits,
            'form': form,
            'doctors': UserProfile.objects.filter(role='dentist'),
            'services': services,
        })

    def _filter_visits(self, form, get_params):
        try:
            params = {}
            if form.is_valid():
                date_from = form.cleaned_data.get('date_from')
                date_to = form.cleaned_data.get('date_to')
                doctor_id = get_params.get('doctor')
                if date_from:
                    params['date_from'] = str(date_from)
                if date_to:
                    params['date_to'] = str(date_to)
                if doctor_id:
                    params['doctor_id'] = doctor_id
            return VisitController.get_all(size=200, **params)
        except Exception:
            return []


# ==================== ПАЦИЕНТ ====================

class PatientDashboardView(PatientRequiredMixin, View):
    template_name = 'patient/dashboard.html'

    def get(self, request):
        patient = PatientHelper.get_patient_for_user(request.user)
        patient_id = patient.get('id') if patient else 0

        try:
            appointments_all = AppointmentController.get_all(patient_id=patient_id, size=200)
            now_str = timezone.now().isoformat()
            upcoming_appointments = [
                a for a in appointments_all
                if a.get('status') == 'scheduled' and str(a.get('datetime', '')) >= now_str
            ]
        except Exception:
            upcoming_appointments = []

        try:
            medical_records = MedicalRecordController.get_all(patient_id=patient_id, size=5)
        except Exception:
            medical_records = []

        try:
            visits = VisitController.get_all(patient_id=patient_id, size=10)
        except Exception:
            visits = []

        try:
            extracts = []
            for v in visits[:5]:
                vid = v.get('id')
                if vid:
                    extracts += VisitController.get_extracts(vid)
        except Exception:
            extracts = []

        context = {
            'patient': patient,
            'upcoming_appointments': upcoming_appointments,
            'visits': visits,
            'medical_records': medical_records,
            'extracts': extracts,
            'stats': self._get_stats(patient_id),
        }
        return render(request, self.template_name, context)

    def _get_stats(self, patient_id):
        try:
            visits = VisitController.get_all(patient_id=patient_id, size=500)
        except Exception:
            visits = []
        return {
            'total_visits': len(visits),
            'total_procedures': 0,
        }


class PatientCancelAppointmentView(PatientRequiredMixin, View):

    def post(self, request, pk):
        patient = PatientHelper.get_patient_for_user(request.user)
        if not patient:
            messages.error(request, 'Профиль пациента не найден')
            return redirect('app:patient_dashboard')

        try:
            appointment = AppointmentController.get_by_id(pk)
        except Exception:
            messages.error(request, 'Запись не найдена')
            return redirect('app:patient_dashboard')

        if appointment.get('patient_id') != patient.get('id'):
            messages.error(request, 'Вы не можете отменить чужую запись')
            return redirect('app:patient_dashboard')

        if appointment.get('status') != 'scheduled':
            messages.warning(request, 'Эта запись уже не может быть отменена')
            return redirect('app:patient_dashboard')

        old_status = appointment.get('status', 'scheduled')
        AppointmentController.update_status(pk, 'cancelled')
        try:
            AppointmentLogController.create({
                'appointment_id': pk,
                'changed_by_login': request.user.username,
                'old_status': old_status,
                'new_status': 'cancelled',
                'comment': f"Отменено пациентом {request.user.get_full_name()}",
            })
        except Exception:
            pass

        messages.success(request, 'Запись отменена')
        return redirect('app:patient_dashboard')


# ==================== ОТЧЁТЫ ====================

class RevenueReportView(ManagerRequiredMixin, View):
    template_name = 'reports/revenue.html'

    def get(self, request):
        period = request.GET.get('period', 'month')
        start_date, end_date = self._get_date_range(period)

        try:
            visits = VisitController.get_all(size=500, date_from=str(start_date), date_to=str(end_date))
        except Exception:
            visits = []

        total_revenue = 0
        visits_count = len(visits)

        context = {
            'total_revenue': total_revenue,
            'top_services': [],
            'doctor_stats': [],
            'visits_count': visits_count,
            'period': period,
            'start_date': start_date,
            'end_date': end_date,
        }
        return render(request, self.template_name, context)

    def _get_date_range(self, period):
        today = timezone.now().date()
        days_map = {'week': 7, 'month': 30, 'quarter': 90, 'year': 365}
        days = days_map.get(period, 30)
        return today - datetime.timedelta(days=days), today


# ==================== УПРАВЛЕНИЕ РОЛЯМИ ====================

class RoleManagerView(AdminRequiredMixin, View):
    template_name = 'admin_panel/role_manager.html'

    def get(self, request):
        users = User.objects.prefetch_related('profile').all()
        role_display = dict(UserProfile.ROLE_CHOICES)

        for user in users:
            profile = getattr(user, 'profile', None)
            user.role = profile.role if profile else 'patient'
            user.get_role_display = role_display.get(user.role, 'Пациент')

        return render(request, self.template_name, {
            'users': users,
            'roles': UserProfile.ROLE_CHOICES
        })


class ChangeRoleView(AdminRequiredMixin, View):

    def post(self, request, user_id):
        new_role = request.POST.get('role')
        user = get_object_or_404(User, id=user_id)

        profile, created = UserProfile.objects.get_or_create(user=user)
        profile.role = new_role
        profile.save()

        self._sync_to_fastapi(user, new_role)
        messages.success(request, f'Роль пользователя {user.username} изменена на {new_role}')
        return redirect('app:role_manager')

    def _sync_to_fastapi(self, user, role):
        try:
            from app.services.fastapi_client import sync_user_to_fastapi
            sync_user_to_fastapi(user, role)
        except (ImportError, Exception) as e:
            print(f"Sync error: {e}")


class CreateUserWithRoleView(AdminRequiredMixin, View):

    def post(self, request):
        username = request.POST.get('username')

        if User.objects.filter(username=username).exists():
            messages.error(request, 'Пользователь с таким логином уже существует')
            return redirect('app:role_manager')

        user = User.objects.create_user(
            username=username,
            email=request.POST.get('email'),
            password=request.POST.get('password'),
            first_name=request.POST.get('first_name', ''),
            last_name=request.POST.get('last_name', '')
        )

        role = request.POST.get('role')
        profile, created = UserProfile.objects.get_or_create(user=user, defaults={'role': role})
        if not created:
            profile.role = role
            profile.save()

        ChangeRoleView()._sync_to_fastapi(user, role)
        messages.success(request, f'Пользователь {username} создан с ролью {role}')
        return redirect('app:role_manager')


class ImpersonateUserView(AdminRequiredMixin, View):

    def get(self, request, user_id):
        user = get_object_or_404(User, id=user_id)
        auth_login(request, user)
        messages.success(request, f'Вы вошли как {user.username}')
        return redirect('app:index')


# ==================== API И ИНТЕГРАЦИЯ ====================

@api_view(['GET'])
@permission_classes([AllowAny])
def api_services(request):
    qs = Service.objects.all().values('id', 'code', 'name', 'cost', 'duration_minutes', 'material_cost')
    if not qs.exists():
        try:
            services = ServiceController.get_all()
        except Exception:
            services = []
    else:
        services = list(qs)
    return Response(services)


class FastAPIDemoView(View):
    template_name = 'integration/fastapi_demo.html'

    def get(self, request):
        return render(request, self.template_name)


class FastAPIServicesView(View):
    template_name = 'fastapi_services.html'

    def get(self, request):
        try:
            import requests
            from django.conf import settings
            base = getattr(settings, 'FASTAPI_URL', 'http://localhost:8000')
            response = requests.get(f"{base}/api/v2/mkbs/services", timeout=5)
            services = response.json() if response.status_code == 200 else []
        except Exception:
            services = []
        return render(request, self.template_name, {'services': services})


class FastAPIServiceStatusView(View):

    def get(self, request):
        status = fastapi_client.get_status_sync()
        return JsonResponse(status)


def fastapi_ping(request):
    """Лёгкий liveness-пробник для фронтенд-поллера.

    Делает живую проверку FastAPI, синхронизирует кэш middleware/баннера
    и возвращает {"online": bool}. Вызывается из base.html и 503_fastapi.html
    каждую секунду — чтобы данные исчезали сразу при выключении API.
    """
    from django.core.cache import cache
    from .services.fastapi_health import get_fastapi_status

    online = get_fastapi_status().get('available', False)
    # держим кэш middleware и баннера в актуальном состоянии
    cache.set('fastapi_available_middleware', online, 5)
    cache.set('fastapi_online_banner', online, 5)
    return JsonResponse({'online': online})


class FastAPIPatientsView(View):

    def get(self, request):
        result = AsyncHelper.run_async(fastapi_client.get_patients())
        return JsonResponse(result)


class FastAPIServicesDataView(View):

    def get(self, request):
        result = AsyncHelper.run_async(fastapi_client.get_services())

        if result['success']:
            return JsonResponse({'success': True, 'data': result['data']})
        return JsonResponse({'success': False, 'error': result['error']}, status=503)


class FastAPIMkbsView(View):

    def get(self, request):
        search = request.GET.get('search', None)
        result = AsyncHelper.run_async(fastapi_client.get_mkbs_codes(search))

        if result['success']:
            return JsonResponse({'success': True, 'data': result['data']})
        return JsonResponse({'success': False, 'error': result['error']}, status=503)


@method_decorator(csrf_exempt, name='dispatch')
class FastAPISyncPatientView(View):

    def post(self, request):
        try:
            data = json.loads(request.body)
            patient_id = data.get('patient_id')

            if not patient_id:
                return JsonResponse({'error': 'patient_id required'}, status=400)

            status = fastapi_client.get_status_sync()
            if not status['online']:
                return JsonResponse({
                    'success': False,
                    'error': 'FastAPI сервер недоступен. Невозможно синхронизировать данные.'
                }, status=503)

            return JsonResponse({
                'success': True,
                'message': f'Пациент {patient_id} синхронизирован',
                'fastapi_status': status
            })

        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)


class FastAPIFullStatusView(View):

    def get(self, request):
        status = fastapi_client.get_status_sync()
        services_result = AsyncHelper.run_async(fastapi_client.get_services())

        return JsonResponse({
            'fastapi': status,
            'services': {
                'available': services_result['success'],
                'count': len(services_result.get('data', [])) if services_result['success'] else 0,
                'error': services_result.get('error')
            },
            'message': 'FastAPI работает нормально' if status['online'] else 'ВНИМАНИЕ: FastAPI сервер НЕ ДОСТУПЕН!'
        })


class FastAPIStatusPageView(View):
    template_name = 'integration/fastapi_status.html'

    def get(self, request):
        return render(request, self.template_name)


# API функции остаются как есть (они уже компактны)
def fastapi_patients(request):
    try:
        import requests
        from django.conf import settings
        base = getattr(settings, 'FASTAPI_URL', 'http://localhost:8000')
        response = requests.get(f"{base}/api/v2/patients/", timeout=5)
        patients = response.json() if response.status_code == 200 else []
    except Exception:
        patients = []
    return JsonResponse({'patients': patients})


def api_patients(request):
    try:
        patients = PatientController.get_all(size=200)
        data = [{'id': p.get('id'), 'full_name': p.get('full_name'), 'phone': p.get('phone')} for p in patients]
    except Exception:
        data = []
    return JsonResponse(data, safe=False)


def api_appointments(request):
    try:
        appointments = AppointmentController.get_all(size=200)
        data = [{'id': a.get('id'), 'patient_id': a.get('patient_id'),
                 'datetime': a.get('datetime'), 'status': a.get('status')} for a in appointments]
    except Exception:
        data = []
    return JsonResponse(data, safe=False)