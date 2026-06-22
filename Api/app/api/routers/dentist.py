from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks, status
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import List, Optional
from datetime import datetime
import tempfile
import os
from docx import Document

from app.core.database import get_db
from app.models.client_models import Clients
from app.models.medical_models import (
    Patient, Appointment, MedicalRecord, Study,
    Referral, WorkOrder, MKBSCode
)
from app.schemas.medical_schemas import (
    AppointmentCreate, AppointmentOut,
    MedicalRecordCreate, MedicalRecordOut,
    StudyCreate, StudyOut,
    ReferralCreate, ReferralOut,
    WorkOrderCreate, WorkOrderOut,
    PatientOut,
    MedicalRecordCreateExtended
)
from app.core.roles import require_dentist
from app.services.mkbs_dll import mkbs_emulator
from app.services.diagnosis_autocomplete import DiagnosisAutocompleteService

router = APIRouter(prefix="/dentist", tags=["Dentist"])


# ========== Запись на прием (через дневник врача) ==========
@router.post("/appointments", response_model=AppointmentOut)
async def create_appointment_by_doctor(
    data: AppointmentCreate,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    patient = await db.get(Patient, data.patient_id)
    if not patient:
        raise HTTPException(404, "Пациент не найден")
    appointment = Appointment(
        patient_id=data.patient_id,
        doctor_id=current_user.id,
        datetime=data.datetime,
        status="scheduled"
    )
    db.add(appointment)
    await db.commit()
    await db.refresh(appointment)
    return appointment


@router.delete("/appointments/{appointment_id}", status_code=status.HTTP_204_NO_CONTENT)
async def cancel_appointment_by_doctor(
    appointment_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    appointment = await db.get(Appointment, appointment_id)
    if not appointment:
        raise HTTPException(404, "Запись не найдена")
    if appointment.doctor_id != current_user.id:
        raise HTTPException(403, "Можно отменять только свои записи")
    appointment.status = "cancelled"
    await db.commit()
    return None


# ========== Дневник врача (расписание на сегодня) ==========
@router.get("/schedule", response_model=List[AppointmentOut])
async def get_doctor_schedule(
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    today_end = today_start.replace(hour=23, minute=59, second=59)
    result = await db.execute(
        select(Appointment).where(
            Appointment.doctor_id == current_user.id,
            Appointment.datetime >= today_start,
            Appointment.datetime <= today_end
        )
    )
    return result.scalars().all()


# ========== Медицинские записи (с поддержкой МКБ-С-3) ==========
@router.post("/medical-records", response_model=MedicalRecordOut)
async def create_medical_record(
    data: MedicalRecordCreateExtended,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    patient = await db.get(Patient, data.patient_id)
    if not patient:
        raise HTTPException(404, "Пациент не найден")

    # Если указан код МКБ-С-3 через ID
    if data.mkbs_code_id:
        code = await db.get(MKBSCode, data.mkbs_code_id)
        if not code:
            raise HTTPException(400, "Неверный код МКБ-С-3")
        if not data.diagnosis:
            data.diagnosis = code.name

    # Если указан строковый код МКБ-С-3
    if hasattr(data, 'mkbs_code_str') and data.mkbs_code_str:
        code_info = mkbs_emulator.get_code_info(data.mkbs_code_str)
        if code_info:
            if not data.diagnosis:
                data.diagnosis = code_info.get("name", "")
            existing = await db.execute(
                select(MKBSCode).where(MKBSCode.code == data.mkbs_code_str)
            )
            db_code = existing.scalar_one_or_none()
            if not db_code:
                db_code = MKBSCode(
                    code=data.mkbs_code_str,
                    name=code_info.get("name", ""),
                    category=code_info.get("category", "diagnosis")
                )
                db.add(db_code)
                await db.flush()
            data.mkbs_code_id = db_code.id

    record_data = data.model_dump(exclude={'mkbs_code_str'}, exclude_unset=True)
    record = MedicalRecord(doctor_id=current_user.id, **record_data)
    db.add(record)
    await db.commit()
    await db.refresh(record)
    return record


@router.get("/medical-records/{patient_id}", response_model=List[MedicalRecordOut])
async def get_patient_medical_records(
    patient_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    result = await db.execute(
        select(MedicalRecord).where(MedicalRecord.patient_id == patient_id)
    )
    return result.scalars().all()


# ========== Поиск по МКБ-С-3 (автозаполнение) ==========
@router.get("/mkbs/search")
async def search_mkbs_codes(
    query: str,
    category: Optional[str] = None,
    current_user: Clients = Depends(require_dentist)
):
    """Поиск кодов МКБ-С-3 для автозаполнения"""
    if category == "diagnosis":
        results = DiagnosisAutocompleteService.search_diagnosis(query)
    elif category == "service":
        results = DiagnosisAutocompleteService.search_services(query)
    else:
        results = DiagnosisAutocompleteService.search_diagnosis(query) + \
                  DiagnosisAutocompleteService.search_services(query)
    return results[:20]


@router.get("/mkbs/validate/{code}")
async def validate_mkbs_code(
    code: str,
    current_user: Clients = Depends(require_dentist)
):
    """Проверить валидность кода МКБ-С-3"""
    is_valid = DiagnosisAutocompleteService.validate_diagnosis_code(code)
    return {"code": code, "valid": is_valid}


# ========== Исследования ==========
@router.post("/studies", response_model=StudyOut)
async def create_study(
    data: StudyCreate,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    study = Study(**data.model_dump())
    db.add(study)
    await db.commit()
    await db.refresh(study)
    return study


@router.get("/studies/{patient_id}", response_model=List[StudyOut])
async def get_patient_studies(
    patient_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    result = await db.execute(select(Study).where(Study.patient_id == patient_id))
    return result.scalars().all()


# ========== Направления ==========
@router.post("/referrals", response_model=ReferralOut)
async def create_referral(
    data: ReferralCreate,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    referral = Referral(doctor_id=current_user.id, **data.model_dump())
    db.add(referral)
    await db.commit()
    await db.refresh(referral)
    return referral


@router.get("/referrals/{patient_id}", response_model=List[ReferralOut])
async def get_patient_referrals(
    patient_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    result = await db.execute(select(Referral).where(Referral.patient_id == patient_id))
    return result.scalars().all()


# ========== Наряды на работу ==========
@router.post("/work-orders", response_model=WorkOrderOut)
async def create_work_order(
    data: WorkOrderCreate,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    work_order = WorkOrder(doctor_id=current_user.id, **data.model_dump())
    db.add(work_order)
    await db.commit()
    await db.refresh(work_order)
    return work_order


@router.get("/work-orders/{patient_id}", response_model=List[WorkOrderOut])
async def get_patient_work_orders(
    patient_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    result = await db.execute(select(WorkOrder).where(WorkOrder.patient_id == patient_id))
    return result.scalars().all()


# ========== Поиск пациентов по дате посещения ==========
@router.get("/patients/search/by-visit-date", response_model=List[PatientOut])
async def search_patients_by_visit_date(
    date: datetime,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    start_date = datetime(date.year, date.month, date.day, 0, 0, 0)
    end_date = datetime(date.year, date.month, date.day, 23, 59, 59)
    subquery = select(MedicalRecord.patient_id).where(
        MedicalRecord.visit_date >= start_date,
        MedicalRecord.visit_date <= end_date
    )
    patient_ids_result = await db.execute(subquery)
    patient_ids = patient_ids_result.scalars().all()
    if not patient_ids:
        return []
    patients_result = await db.execute(select(Patient).where(Patient.id.in_(patient_ids)))
    return patients_result.scalars().all()


# ========== Выписка (Word) ==========
@router.post("/discharge/{patient_id}")
async def create_discharge(
    patient_id: int,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: Clients = Depends(require_dentist)
):
    patient = await db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(404, "Пациент не найден")
    records_result = await db.execute(
        select(MedicalRecord)
        .where(MedicalRecord.patient_id == patient_id)
        .order_by(MedicalRecord.visit_date.desc())
    )
    records = records_result.scalars().all()
    if not records:
        raise HTTPException(404, "Нет медицинских записей")

    doc = Document()
    doc.add_heading('Медицинская карта стоматологического больного (форма 043/у)', 0)
    doc.add_paragraph(f"ФИО пациента: {patient.full_name}")
    doc.add_paragraph(f"Дата рождения: {patient.birth_date.strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Телефон: {patient.phone}")
    doc.add_heading('Посещения:', level=1)
    for rec in records:
        doc.add_heading(f"Дата: {rec.visit_date.strftime('%d.%m.%Y %H:%M')}", level=2)
        doc.add_paragraph(f"Жалобы: {rec.complaints or '—'}")
        doc.add_paragraph(f"Анамнез: {rec.anamnesis or '—'}")
        doc.add_paragraph(f"Осмотр: {rec.examination or '—'}")
        doc.add_paragraph(f"Диагноз: {rec.diagnosis or '—'}")
        if rec.mkbs_code:
            doc.add_paragraph(f"Код МКБ-С-3: {rec.mkbs_code.code} - {rec.mkbs_code.name}")
        doc.add_paragraph(f"Назначения: {rec.prescriptions or '—'}")
        doc.add_paragraph(f"Зубная формула: {rec.tooth_formula or '—'}")
        doc.add_paragraph(f"Аллергии: {rec.alert_info or '—'}")
        doc.add_paragraph("-" * 40)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    background_tasks.add_task(os.unlink, tmp_path)
    return FileResponse(
        tmp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"discharge_{patient_id}.docx"
    )