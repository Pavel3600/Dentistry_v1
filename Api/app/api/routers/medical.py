from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import select
from typing import List
from datetime import datetime
import tempfile
import os

from app.db.session import get_session
from app.models.client_models import Clients
from app.models.medical_models import Patient, Appointment, MedicalRecord, Study, Referral, WorkOrder
from app.schemas.medical_schemas import (
    PatientCreate, PatientUpdate, PatientOut,
    AppointmentCreate, AppointmentOut,
    MedicalRecordCreate, MedicalRecordOut,
    StudyCreate, StudyOut,
    ReferralCreate, ReferralOut,
    WorkOrderCreate, WorkOrderOut
)
from app.auth.roles import require_manager, require_dentist, require_admin

router = APIRouter()

# ---------- Менеджер и стоматолог: работа с пациентами ----------
@router.get("/patients", response_model=List[PatientOut], tags=["Medical"])
def get_all_patients(
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_manager)
):
    if current_user.role not in ["manager", "dentist", "admin"]:
        raise HTTPException(status_code=403, detail="Недостаточно прав")
    patients = session.execute(select(Patient)).scalars().all()
    return patients

@router.get("/patients/{patient_id}", response_model=PatientOut, tags=["Medical"])
def get_patient(
    patient_id: int,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_manager)
):
    patient = session.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Пациент не найден")
    return patient

@router.post("/patients", response_model=PatientOut, status_code=status.HTTP_201_CREATED, tags=["Medical"])
def create_patient(
    data: PatientCreate,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_manager)
):
    user = session.get(Clients, data.user_id)
    if not user or user.role != "patient":
        raise HTTPException(status_code=400, detail="Неверный user_id или пользователь не пациент")
    existing = session.execute(select(Patient).where(Patient.user_id == data.user_id)).first()
    if existing:
        raise HTTPException(status_code=400, detail="Пациент уже привязан к этому пользователю")
    patient = Patient(**data.model_dump())
    session.add(patient)
    session.commit()
    session.refresh(patient)
    return patient

@router.put("/patients/{patient_id}", response_model=PatientOut, tags=["Medical"])
def update_patient(
    patient_id: int,
    data: PatientUpdate,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_manager)
):
    patient = session.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Пациент не найден")
    for key, value in data.model_dump(exclude_unset=True).items():
        setattr(patient, key, value)
    session.commit()
    session.refresh(patient)
    return patient

@router.delete("/patients/{patient_id}", tags=["Medical"])
def delete_patient(
    patient_id: int,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_manager)
):
    patient = session.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Пациент не найден")
    session.delete(patient)
    session.commit()
    return {"ok": True}

# ---------- Запись на приём (менеджер) ----------
@router.post("/appointments", response_model=AppointmentOut, tags=["Medical"])
def create_appointment(
    data: AppointmentCreate,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_manager)
):
    patient = session.get(Patient, data.patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Пациент не найден")
    doctor = session.get(Clients, data.doctor_id)
    if not doctor or doctor.role != "dentist":
        raise HTTPException(status_code=400, detail="Врач не является стоматологом")
    appointment = Appointment(**data.model_dump())
    session.add(appointment)
    session.commit()
    session.refresh(appointment)
    return appointment

@router.get("/appointments", response_model=List[AppointmentOut], tags=["Medical"])
def get_appointments(
    patient_id: int = None,
    doctor_id: int = None,
    date_from: datetime = None,
    date_to: datetime = None,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    query = select(Appointment)
    if patient_id:
        query = query.where(Appointment.patient_id == patient_id)
    if doctor_id:
        query = query.where(Appointment.doctor_id == doctor_id)
    if date_from:
        query = query.where(Appointment.datetime >= date_from)
    if date_to:
        query = query.where(Appointment.datetime <= date_to)
    appointments = session.execute(query).scalars().all()
    return appointments

@router.delete("/appointments/{appointment_id}", tags=["Medical"])
def cancel_appointment(
    appointment_id: int,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_manager)
):
    appointment = session.get(Appointment, appointment_id)
    if not appointment:
        raise HTTPException(status_code=404, detail="Запись не найдена")
    appointment.status = "cancelled"
    session.commit()
    return {"message": "Запись отменена"}

# ---------- Запись на приём для стоматолога (через дневник) ----------
@router.post("/doctor/appointments", response_model=AppointmentOut, tags=["Medical"])
def create_appointment_by_doctor(
    data: AppointmentCreate,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    patient = session.get(Patient, data.patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Пациент не найден")
    appointment = Appointment(
        patient_id=data.patient_id,
        doctor_id=current_user.id,
        datetime=data.datetime,
        status="scheduled"
    )
    session.add(appointment)
    session.commit()
    session.refresh(appointment)
    return appointment

@router.delete("/doctor/appointments/{appointment_id}", tags=["Medical"])
def cancel_appointment_by_doctor(
    appointment_id: int,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    appointment = session.get(Appointment, appointment_id)
    if not appointment:
        raise HTTPException(status_code=404, detail="Запись не найдена")
    if appointment.doctor_id != current_user.id:
        raise HTTPException(status_code=403, detail="Можно отменять только свои записи")
    appointment.status = "cancelled"
    session.commit()
    return {"message": "Запись отменена"}

# ---------- Медицинские карты (только стоматолог) ----------
@router.post("/medical-records", response_model=MedicalRecordOut, tags=["Medical"])
def create_medical_record(
    data: MedicalRecordCreate,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    patient = session.get(Patient, data.patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Пациент не найден")
    record = MedicalRecord(doctor_id=current_user.id, **data.model_dump())
    session.add(record)
    session.commit()
    session.refresh(record)
    return record

@router.get("/medical-records/{patient_id}", response_model=List[MedicalRecordOut], tags=["Medical"])
def get_patient_medical_records(
    patient_id: int,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    records = session.execute(select(MedicalRecord).where(MedicalRecord.patient_id == patient_id)).scalars().all()
    return records

# ---------- Исследования (стоматолог) ----------
@router.post("/studies", response_model=StudyOut, tags=["Medical"])
def create_study(
    data: StudyCreate,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    study = Study(**data.model_dump())
    session.add(study)
    session.commit()
    session.refresh(study)
    return study

@router.get("/studies/{patient_id}", response_model=List[StudyOut], tags=["Medical"])
def get_patient_studies(
    patient_id: int,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    studies = session.execute(select(Study).where(Study.patient_id == patient_id)).scalars().all()
    return studies

# ---------- Направления (стоматолог) ----------
@router.post("/referrals", response_model=ReferralOut, tags=["Medical"])
def create_referral(
    data: ReferralCreate,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    referral = Referral(doctor_id=current_user.id, **data.model_dump())
    session.add(referral)
    session.commit()
    session.refresh(referral)
    return referral

@router.get("/referrals/{patient_id}", response_model=List[ReferralOut], tags=["Medical"])
def get_patient_referrals(
    patient_id: int,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    referrals = session.execute(select(Referral).where(Referral.patient_id == patient_id)).scalars().all()
    return referrals

# ---------- Наряды на работу (стоматолог) ----------
@router.post("/work-orders", response_model=WorkOrderOut, tags=["Medical"])
def create_work_order(
    data: WorkOrderCreate,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    work_order = WorkOrder(doctor_id=current_user.id, **data.model_dump())
    session.add(work_order)
    session.commit()
    session.refresh(work_order)
    return work_order

@router.get("/work-orders/{patient_id}", response_model=List[WorkOrderOut], tags=["Medical"])
def get_patient_work_orders(
    patient_id: int,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    orders = session.execute(select(WorkOrder).where(WorkOrder.patient_id == patient_id)).scalars().all()
    return orders

# ---------- Поиск пациентов по дате посещения ----------
@router.get("/patients/search/by-visit-date", response_model=List[PatientOut], tags=["Medical"])
def search_patients_by_visit_date(
    date: datetime,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    start_date = datetime(date.year, date.month, date.day, 0, 0, 0)
    end_date = datetime(date.year, date.month, date.day, 23, 59, 59)
    subquery = select(MedicalRecord.patient_id).where(
        MedicalRecord.visit_date >= start_date,
        MedicalRecord.visit_date <= end_date
    )
    patient_ids = session.execute(subquery).scalars().all()
    patients = session.execute(select(Patient).where(Patient.id.in_(patient_ids))).scalars().all()
    return patients

# ---------- Дневник врача (расписание приёмов на сегодня) ----------
@router.get("/doctor/schedule", response_model=List[AppointmentOut], tags=["Medical"])
def get_doctor_schedule(
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    today_end = today_start.replace(hour=23, minute=59, second=59)
    appointments = session.execute(
        select(Appointment).where(
            Appointment.doctor_id == current_user.id,
            Appointment.datetime >= today_start,
            Appointment.datetime <= today_end
        )
    ).scalars().all()
    return appointments

# ---------- Формирование выписки (Word) ----------
@router.post("/discharge/{patient_id}", tags=["Medical"])
def create_discharge(
    patient_id: int,
    background_tasks: BackgroundTasks,
    session: Session = Depends(get_session),
    current_user: Clients = Depends(require_dentist)
):
    from docx import Document
    patient = session.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Пациент не найден")

    records = session.execute(
        select(MedicalRecord).where(MedicalRecord.patient_id == patient_id).order_by(MedicalRecord.visit_date.desc())
    ).scalars().all()

    if not records:
        raise HTTPException(status_code=404, detail="Нет медицинских записей")

    doc = Document()
    doc.add_heading('Медицинская карта стоматологического больного (форма 043/у)', 0)
    doc.add_paragraph(f"ФИО пациента: {patient.full_name}")
    doc.add_paragraph(f"Дата рождения: {patient.birth_date.strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Телефон: {patient.phone}")
    doc.add_heading('Посещения:', level=1)

    for rec in records:
        doc.add_heading(f"Дата: {rec.visit_date.strftime('%d.%m.%Y %H:%M')}", level=2)
        doc.add_paragraph(f"Жалобы: {rec.complaints or '—'}")
        doc.add_paragraph(f"Анамнез: {rec.anamnesis or '—'}")
        doc.add_paragraph(f"Осмотр: {rec.examination or '—'}")
        doc.add_paragraph(f"Диагноз: {rec.diagnosis or '—'}")
        doc.add_paragraph(f"Назначения: {rec.prescriptions or '—'}")
        doc.add_paragraph(f"Зубная формула: {rec.tooth_formula or '—'}")
        doc.add_paragraph(f"Аллергии: {rec.alert_info or '—'}")
        doc.add_paragraph("-" * 40)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    background_tasks.add_task(os.unlink, tmp_path)

    return FileResponse(
        tmp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"discharge_{patient_id}.docx"
    )